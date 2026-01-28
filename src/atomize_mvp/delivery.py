from pathlib import Path

from docx import Document

from atomize_mvp.schemas import DraftsSchema


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def write_linkedin_docx(path: Path, drafts: DraftsSchema) -> None:
    doc = Document()
    doc.add_heading("LinkedIn Posts", level=1)
    for post in drafts.linkedin_posts:
        doc.add_heading(post.id, level=2)
        doc.add_paragraph(post.hook)
        doc.add_paragraph(post.body)
        doc.add_paragraph(f"CTA: {post.cta}")
        if post.hashtags:
            doc.add_paragraph("Hashtags: " + " ".join(post.hashtags))
    _ensure_parent(path)
    doc.save(path)


def write_x_threads_docx(path: Path, drafts: DraftsSchema) -> None:
    doc = Document()
    doc.add_heading("X Threads", level=1)
    for thread in drafts.x_threads:
        doc.add_heading(thread.id, level=2)
        for idx, tweet in enumerate(thread.tweets, start=1):
            doc.add_paragraph(f"{idx}. {tweet}")
        doc.add_paragraph(f"Closing CTA: {thread.closing_cta}")
    _ensure_parent(path)
    doc.save(path)


def write_blog_outlines_docx(path: Path, drafts: DraftsSchema) -> None:
    doc = Document()
    doc.add_heading("Blog Outlines", level=1)
    for outline in drafts.blog_outlines:
        doc.add_heading(outline.id, level=2)
        doc.add_paragraph(f"Title: {outline.title}")
        doc.add_paragraph(f"Audience: {outline.audience}")
        doc.add_paragraph(f"Goal: {outline.goal}")
        doc.add_paragraph("Outline:")
        for item in outline.outline:
            doc.add_paragraph(f"- {item}")
        doc.add_paragraph("Key takeaways:")
        for item in outline.key_takeaways:
            doc.add_paragraph(f"- {item}")
    _ensure_parent(path)
    doc.save(path)


def write_ig_stories_docx(path: Path, drafts: DraftsSchema) -> None:
    doc = Document()
    doc.add_heading("IG Stories", level=1)
    for story in drafts.ig_stories:
        doc.add_heading(story.id, level=2)
        for idx, slide in enumerate(story.slides, start=1):
            doc.add_paragraph(f"Slide {idx}: {slide}")
    _ensure_parent(path)
    doc.save(path)
