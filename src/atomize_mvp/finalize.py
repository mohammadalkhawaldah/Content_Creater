from __future__ import annotations

import csv
import shutil
from dataclasses import dataclass
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from atomize_mvp.schemas import ContentBlueprint, DraftsSchema


@dataclass
class FinalizeOutputs:
    content_library: Path
    hooks_quotes: Path
    readme: Path
    schedule: Path | None


def _align_paragraph(paragraph, lang: str) -> None:
    if lang == "ar":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_heading(doc: Document, text: str, level: int, lang: str) -> None:
    para = doc.add_heading(text, level=level)
    _align_paragraph(para, lang)


def _add_paragraph(doc: Document, text: str, lang: str, style: str | None = None) -> None:
    para = doc.add_paragraph(text, style=style)
    _align_paragraph(para, lang)


def _write_content_library(
    path: Path,
    client: str,
    title: str,
    lang: str,
    drafts: DraftsSchema,
) -> None:
    doc = Document()
    _add_heading(doc, "Atomize Content Library", 0, lang)
    _add_paragraph(doc, f"Client: {client}", lang)
    _add_paragraph(doc, f"Title: {title}", lang)
    _add_paragraph(doc, f"Date: {date.today().isoformat()}", lang)
    _add_paragraph(
        doc,
        "Counts: "
        f"LinkedIn {len(drafts.linkedin_posts)}, "
        f"X {len(drafts.x_threads)}, "
        f"Blog {len(drafts.blog_outlines)}, "
        f"IG {len(drafts.ig_stories)}",
        lang,
    )

    _add_heading(doc, "LinkedIn Posts", 1, lang)
    for post in drafts.linkedin_posts:
        _add_heading(doc, post.id, 2, lang)
        _add_paragraph(doc, post.hook, lang)
        _add_paragraph(doc, post.body, lang)
        _add_paragraph(doc, f"CTA: {post.cta}", lang)
        if post.hashtags:
            _add_paragraph(doc, "Hashtags: " + " ".join(post.hashtags), lang)

    _add_heading(doc, "X Threads", 1, lang)
    for thread in drafts.x_threads:
        _add_heading(doc, thread.id, 2, lang)
        for idx, tweet in enumerate(thread.tweets, start=1):
            _add_paragraph(doc, f"{idx}. {tweet}", lang)
        _add_paragraph(doc, f"Closing CTA: {thread.closing_cta}", lang)

    _add_heading(doc, "Blog Outlines", 1, lang)
    for outline in drafts.blog_outlines:
        _add_heading(doc, outline.id, 2, lang)
        _add_paragraph(doc, f"Title: {outline.title}", lang)
        _add_paragraph(doc, f"Audience: {outline.audience}", lang)
        _add_paragraph(doc, f"Goal: {outline.goal}", lang)
        _add_paragraph(doc, "Outline:", lang)
        for item in outline.outline:
            _add_paragraph(doc, f"- {item}", lang)
        _add_paragraph(doc, "Key takeaways:", lang)
        for item in outline.key_takeaways:
            _add_paragraph(doc, f"- {item}", lang)

    _add_heading(doc, "Instagram Stories", 1, lang)
    for story in drafts.ig_stories:
        _add_heading(doc, story.id, 2, lang)
        for idx, slide in enumerate(story.slides, start=1):
            _add_paragraph(doc, f"Slide {idx}: {slide}", lang)

    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx(doc, path)


def _write_hooks_quotes(path: Path, lang: str, blueprint: ContentBlueprint) -> None:
    doc = Document()
    _add_heading(doc, "Hooks & Quotes", 0, lang)

    _add_heading(doc, "Hooks", 1, lang)
    for hook in blueprint.hooks:
        _add_paragraph(doc, hook, lang, style="List Bullet")

    _add_heading(doc, "Quotes", 1, lang)
    for quote in blueprint.quotes:
        _add_paragraph(doc, quote, lang, style="List Bullet")

    _add_heading(doc, "CTAs", 1, lang)
    for cta in blueprint.ctas:
        _add_paragraph(doc, cta, lang, style="List Bullet")

    _add_heading(doc, "Key Points", 1, lang)
    for point in blueprint.key_points:
        _add_paragraph(doc, point, lang, style="List Bullet")

    _add_heading(doc, "Do Not Say", 1, lang)
    for phrase in blueprint.do_not_say:
        _add_paragraph(doc, phrase, lang, style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx(doc, path)


def _write_readme(
    path: Path,
    client: str,
    title: str,
    lang: str,
    tone: str,
    drafts: DraftsSchema,
) -> None:
    doc = Document()
    _add_heading(doc, f"README - Start Here - {client} - {title}", 0, lang)
    _add_paragraph(doc, "What is included in this delivery:", lang)
    _add_paragraph(doc, "- Content Library (all platform drafts)", lang)
    _add_paragraph(doc, "- Hooks & Quotes reference", lang)
    _add_paragraph(doc, "- Platform-ready DOCX files", lang)
    _add_paragraph(doc, "- Transcripts and segments", lang)
    _add_paragraph(doc, "How to use:", lang)
    _add_paragraph(doc, "- Copy/paste drafts into your scheduling tools.", lang)
    _add_paragraph(doc, "- Review tone and language before publishing.", lang)
    _add_paragraph(
        doc,
        f"Counts: LinkedIn {len(drafts.linkedin_posts)}, X {len(drafts.x_threads)}, "
        f"Blog {len(drafts.blog_outlines)}, IG {len(drafts.ig_stories)}",
        lang,
    )
    _add_paragraph(doc, f"Tone: {tone}", lang)
    _add_paragraph(doc, f"Language: {lang}", lang)
    _add_paragraph(doc, "Support: contact your Atomize account manager.", lang)
    path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx(doc, path)


def _save_docx(doc: Document, path: Path) -> None:
    temp_path = path.with_suffix(path.suffix + ".tmp")
    try:
        if temp_path.exists():
            temp_path.unlink()
        doc.save(temp_path)
        if path.exists():
            path.unlink()
        temp_path.replace(path)
    except PermissionError as exc:
        raise PermissionError(
            f"Permission denied writing {path}. Close the file if it is open and retry."
        ) from exc


def _build_schedule_rows(
    drafts: DraftsSchema, base_date: date, days: int
) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    end_date = base_date + timedelta(days=days - 1)
    dates = [base_date + timedelta(days=i) for i in range(days)]

    def add_rows(platform: str, items: list, get_title) -> None:
        idx = 0
        for day in dates:
            if idx >= len(items):
                break
            if platform == "LinkedIn" and day.weekday() >= 5:
                continue
            if platform == "X" and day.weekday() not in (0, 2, 4):
                continue
            if platform == "IG" and day.weekday() not in (1, 3):
                continue
            if platform == "Blog" and day.weekday() != 4:
                continue
            item = items[idx]
            rows.append(
                {
                    "date": day.isoformat(),
                    "platform": platform,
                    "content_id": item.id,
                    "title_or_hook": get_title(item),
                    "status": "planned",
                }
            )
            idx += 1

        if idx < len(items):
            for day in dates:
                if idx >= len(items):
                    break
                rows.append(
                    {
                        "date": day.isoformat(),
                        "platform": platform,
                        "content_id": items[idx].id,
                        "title_or_hook": get_title(items[idx]),
                        "status": "planned",
                    }
                )
                idx += 1

    add_rows("LinkedIn", drafts.linkedin_posts, lambda p: p.hook)
    add_rows("X", drafts.x_threads, lambda t: t.tweets[0] if t.tweets else "")
    add_rows("IG", drafts.ig_stories, lambda s: s.slides[0] if s.slides else "")
    add_rows("Blog", drafts.blog_outlines, lambda b: b.title)

    rows = [row for row in rows if base_date <= date.fromisoformat(row["date"]) <= end_date]
    return rows


def _write_schedule_csv(path: Path, drafts: DraftsSchema, base_date: date | None = None) -> Path:
    base = base_date or date.today()
    rows = _build_schedule_rows(drafts, base, 14)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle, fieldnames=["date", "platform", "content_id", "title_or_hook", "status"]
        )
        writer.writeheader()
        writer.writerows(rows)
    return path


def _copy_if_exists(source: Path, target_dir: Path) -> None:
    if source.exists():
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / source.name
        try:
            if target.exists():
                target.unlink()
            shutil.copy2(source, target)
        except PermissionError as exc:
            raise PermissionError(
                f"Permission denied writing {target}. Close the file if it is open and retry."
            ) from exc


def finalize_delivery(
    job_root: Path,
    drafts: DraftsSchema,
    blueprint: ContentBlueprint,
    client: str,
    title: str,
    lang: str,
    tone: str,
    include_schedule: bool = True,
    base_date: date | None = None,
) -> FinalizeOutputs:
    delivery_root = job_root / "04_delivery"
    delivery_root.mkdir(parents=True, exist_ok=True)

    content_library = delivery_root / f"Content Library - {client} - {title}.docx"
    hooks_quotes = delivery_root / f"Hooks & Quotes - {client} - {title}.docx"
    readme = delivery_root / f"README - Start Here - {client} - {title}.docx"
    schedule = delivery_root / "Posting Schedule - 14 Days.csv"

    _write_content_library(content_library, client, title, lang, drafts)
    _write_hooks_quotes(hooks_quotes, lang, blueprint)
    _write_readme(readme, client, title, lang, tone, drafts)

    schedule_path = None
    if include_schedule:
        schedule_path = _write_schedule_csv(schedule, drafts, base_date=base_date)

    final_root = delivery_root / "Final Delivery"
    (final_root / "00_README").mkdir(parents=True, exist_ok=True)
    (final_root / "01_Content_Library").mkdir(parents=True, exist_ok=True)
    (final_root / "02_Hooks_Quotes").mkdir(parents=True, exist_ok=True)
    (final_root / "03_Transcript").mkdir(parents=True, exist_ok=True)
    (final_root / "04_Platform_Ready").mkdir(parents=True, exist_ok=True)
    (final_root / "07_Schedule").mkdir(parents=True, exist_ok=True)

    _copy_if_exists(readme, final_root / "00_README")
    _copy_if_exists(content_library, final_root / "01_Content_Library")
    _copy_if_exists(hooks_quotes, final_root / "02_Hooks_Quotes")

    transcripts_dir = job_root / "02_transcripts"
    for name in ["transcript.txt", "clean_transcript.txt", "transcript.srt", "segments.json"]:
        _copy_if_exists(transcripts_dir / name, final_root / "03_Transcript")

    platform_ready = delivery_root / "Platform Ready"
    if platform_ready.exists():
        for docx in platform_ready.glob("*.docx"):
            _copy_if_exists(docx, final_root / "04_Platform_Ready")

    if schedule_path:
        _copy_if_exists(schedule_path, final_root / "07_Schedule")

    return FinalizeOutputs(
        content_library=content_library,
        hooks_quotes=hooks_quotes,
        readme=readme,
        schedule=schedule_path,
    )
